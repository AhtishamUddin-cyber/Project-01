"""
Core analysis engine for the AI Smart Trade Analyzer web app.
This is the same logic from the original Colab notebook, refactored so
it returns structured data instead of printing to a console — so it
can be rendered in a Streamlit web page instead.
"""

import io
import os
import re
import json
import time
import base64
from datetime import datetime, timedelta

import requests
from PIL import Image
from textblob import TextBlob
from google import genai

GEMINI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash-001",
    "gemini-2.0-flash-lite",
    "gemini-flash-latest",
]

COIN_MAP = {
    "BTC": "bitcoin", "ETH": "ethereum", "SOL": "solana", "BNB": "binancecoin",
    "XRP": "ripple", "ADA": "cardano", "DOGE": "dogecoin", "MATIC": "matic-network",
    "DOT": "polkadot", "LTC": "litecoin", "AVAX": "avalanche-2", "LINK": "chainlink",
    "UNI": "uniswap", "ATOM": "cosmos", "APT": "aptos", "ARB": "arbitrum",
    "OP": "optimism", "SUI": "sui", "TRX": "tron", "TON": "the-open-network",
    "NEAR": "near", "INJ": "injective-protocol", "FTM": "fantom", "PEPE": "pepe",
    "SHIB": "shiba-inu", "WIF": "dogwifcoin", "JUP": "jupiter-exchange-solana",
    "PYTH": "pyth-network",
}

# Bitget lists several futures categories beyond plain crypto perpetuals
# (stock-linked futures like AAPL/TSLA, USDC-margined, coin-margined, etc).
# We try them in this order until one responds with real data.
FUTURES_PRODUCT_TYPES = ["usdt-futures", "susdt-futures", "usdc-futures", "coin-futures"]

PATTERN_LIBRARY_FILE = "pattern_library.json"
TRADE_TRACKER_FILE = "trades.json"


# ─────────────────────────────────────────────────────────────
#   GITHUB-BACKED PERSISTENCE (fixes patterns/trades disappearing)
# ─────────────────────────────────────────────────────────────
# Streamlit Cloud's filesystem is EPHEMERAL: local files like
# pattern_library.json / trades.json get wiped whenever the app container
# restarts. This happens on every redeploy AND whenever a free-tier app
# goes to sleep from inactivity and then wakes back up - nothing was
# actually "deleting" the data on refresh, the whole container (and its
# local disk) was being replaced underneath it.
#
# Fix: also persist these files inside the GitHub repo itself via the
# GitHub Contents API. The repo survives every restart/redeploy (it's what
# triggers the redeploy), so this makes the data genuinely permanent.
# Needs a free GitHub Personal Access Token — falls back to local-file-only
# behavior (old behavior, still works within a session) if no token is set.
GITHUB_REPO = "AhtishamUddin-cyber/Project-01"
GITHUB_BRANCH = "main"


def _github_get(path, token):
    if not token:
        return None, None
    try:
        r = requests.get(
            f"https://api.github.com/repos/{GITHUB_REPO}/contents/{path}",
            headers={"Authorization": f"token {token}", "Accept": "application/vnd.github+json"},
            params={"ref": GITHUB_BRANCH}, timeout=10,
        )
        if r.status_code != 200:
            return None, None
        data = r.json()
        content = base64.b64decode(data["content"]).decode("utf-8")
        return content, data["sha"]
    except Exception:
        return None, None


def _github_put(path, token, content_str, message):
    if not token:
        return False
    try:
        _, sha = _github_get(path, token)
        payload = {
            "message": message,
            "content": base64.b64encode(content_str.encode("utf-8")).decode("utf-8"),
            "branch": GITHUB_BRANCH,
        }
        if sha:
            payload["sha"] = sha
        r = requests.put(
            f"https://api.github.com/repos/{GITHUB_REPO}/contents/{path}",
            headers={"Authorization": f"token {token}", "Accept": "application/vnd.github+json"},
            json=payload, timeout=15,
        )
        return r.status_code in (200, 201)
    except Exception:
        return False


# ─────────────────────────────────────────────────────────────
#   TRADE TRACKER — log real trades, auto-check TP/SL against live price
# ─────────────────────────────────────────────────────────────
def load_trades(github_token=None):
    if github_token:
        content, _ = _github_get("data/trades.json", github_token)
        if content is not None:
            try:
                trades = json.loads(content)
                with open(TRADE_TRACKER_FILE, "w") as f:
                    f.write(content)
                return trades
            except Exception:
                pass
    if os.path.exists(TRADE_TRACKER_FILE):
        with open(TRADE_TRACKER_FILE, "r") as f:
            return json.load(f)
    return []


def save_trades(trades, github_token=None):
    content = json.dumps(trades, indent=2)
    with open(TRADE_TRACKER_FILE, "w") as f:
        f.write(content)
    if github_token:
        _github_put("data/trades.json", github_token, content, "Update trade tracker")


def add_trade(coin_symbol, pair, market_type, direction, entry, tp1, tp2, sl, timeframe, note="",
              github_token=None, confidence=None, vote_margin=None):
    trades = load_trades(github_token)
    trade = {
        "id": f"{pair}_{int(time.time()*1000)}",
        "coin": coin_symbol, "pair": pair, "market_type": market_type,
        "direction": direction, "entry": entry, "tp1": tp1, "tp2": tp2, "sl": sl,
        "timeframe": timeframe, "note": note,
        "status": "OPEN",
        "opened_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "closed_at": None,
        "exit_price": None,
        # Stored so we can later measure REAL win rate by signal strength,
        # instead of just trusting the confidence score in the abstract.
        "confidence": confidence,
        "vote_margin": vote_margin,
    }
    trades.append(trade)
    save_trades(trades, github_token)
    return trade


def position_size(account_balance, risk_pct, entry, sl, leverage=1):
    """How many coins/contracts to buy so that hitting the SL loses exactly
    `risk_pct`% of the account - not more, not less. This is the single
    biggest thing separating a controlled trade from a gambling one: the
    signal decides direction, but position size decides how much damage a
    wrong call can do."""
    if not entry or not sl or entry == sl or not account_balance or not risk_pct:
        return None
    risk_amount = account_balance * (risk_pct / 100)
    price_risk_per_unit = abs(entry - sl)
    units = risk_amount / price_risk_per_unit
    position_value = units * entry
    margin_required = position_value / max(leverage, 1)
    return {
        "units": units,
        "position_value": position_value,
        "margin_required": margin_required,
        "risk_amount": risk_amount,
        "leverage": leverage,
    }


def performance_by_confidence(trades):
    """Real win-rate feedback from the trader's OWN closed trades, bucketed
    by how confident the tool was when the trade was opened. This is the
    honest way to know if the confidence score means anything - not by
    trusting the formula, but by checking what actually happened."""
    closed = [t for t in trades if t["status"] != "OPEN" and t.get("confidence") is not None]
    buckets = {
        "High (75%+)": [t for t in closed if t["confidence"] >= 75],
        "Medium (55-74%)": [t for t in closed if 55 <= t["confidence"] < 75],
        "Low (<55%)": [t for t in closed if t["confidence"] < 55],
    }
    out = []
    for label, group in buckets.items():
        if not group:
            out.append({"bucket": label, "trades": 0, "win_rate": None, "avg_pnl": None})
            continue
        wins = [t for t in group if t["status"] in ("TP1_HIT", "TP2_HIT")]
        avg_pnl = sum(t.get("pnl_pct", 0) or 0 for t in group) / len(group)
        out.append({
            "bucket": label, "trades": len(group),
            "win_rate": round(len(wins) / len(group) * 100, 1),
            "avg_pnl": round(avg_pnl, 2),
        })
    return out


def get_single_ticker_price(pair, market_type="spot"):
    """Fetch just the live price for one symbol (cheap, single-symbol call)."""
    try:
        if market_type == "futures":
            for ptype in FUTURES_PRODUCT_TYPES:
                try:
                    r = requests.get(
                        "https://api.bitget.com/api/v2/mix/market/ticker",
                        params={"symbol": pair, "productType": ptype}, timeout=8,
                    )
                    data = r.json().get("data", [])
                    if data:
                        d = data[0]
                        price = float(d.get("lastPr") or d.get("last") or 0)
                        if price:
                            return price
                except Exception:
                    continue
            return None
        else:
            r = requests.get(
                "https://api.bitget.com/api/v2/spot/market/tickers",
                params={"symbol": pair}, timeout=8,
            )
            data = r.json().get("data", [])
            if not data:
                return None
            d = data[0]
            return float(d.get("lastPr") or d.get("last") or 0) or None
    except Exception:
        return None


def evaluate_trade(trade, live_price):
    """Checks a trade's live price against its TP/SL and returns an updated
    copy with status + unrealized/realized P&L. Doesn't mutate the original."""
    t = dict(trade)
    if live_price is None:
        t["current_price"] = None
        t["pnl_pct"] = None
        return t

    t["current_price"] = live_price
    direction = t["direction"]
    entry = t["entry"]

    if direction == "LONG":
        pnl_pct = ((live_price - entry) / entry) * 100 if entry else 0
    else:
        pnl_pct = ((entry - live_price) / entry) * 100 if entry else 0
    t["pnl_pct"] = round(pnl_pct, 2)

    if t["status"] == "OPEN":
        sl = t.get("sl")
        tp1 = t.get("tp1")
        tp2 = t.get("tp2")

        if direction == "LONG":
            hit_sl = sl and live_price <= sl
            hit_tp2 = tp2 and live_price >= tp2
            hit_tp1 = tp1 and live_price >= tp1
        else:
            hit_sl = sl and live_price >= sl
            hit_tp2 = tp2 and live_price <= tp2
            hit_tp1 = tp1 and live_price <= tp1

        if hit_sl:
            t["status"] = "SL_HIT"
            t["closed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            t["exit_price"] = sl
        elif hit_tp2:
            t["status"] = "TP2_HIT"
            t["closed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            t["exit_price"] = tp2
        elif hit_tp1:
            t["status"] = "TP1_HIT"
            t["closed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            t["exit_price"] = tp1

    return t


def refresh_all_trades(github_token=None):
    """Re-checks every OPEN trade's live price and updates status in storage.
    Returns the fully refreshed list (open + closed)."""
    trades = load_trades(github_token)
    updated = []
    for t in trades:
        if t["status"] == "OPEN":
            price = get_single_ticker_price(t["pair"], t["market_type"])
            t = evaluate_trade(t, price)
        else:
            price = get_single_ticker_price(t["pair"], t["market_type"])
            t = evaluate_trade(t, price)
        updated.append(t)
    save_trades(updated, github_token)
    return updated


def close_trade_manually(trade_id, exit_price=None, note="", github_token=None):
    trades = load_trades(github_token)
    for t in trades:
        if t["id"] == trade_id and t["status"] == "OPEN":
            t["status"] = "CLOSED_MANUAL"
            t["closed_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            t["exit_price"] = exit_price
            if note:
                t["note"] = (t.get("note", "") + " | " + note).strip(" |")
    save_trades(trades, github_token)
    return trades


def delete_trade(trade_id, github_token=None):
    trades = load_trades(github_token)
    trades = [t for t in trades if t["id"] != trade_id]
    save_trades(trades, github_token)
    return trades


def trade_stats(trades):
    """Win-rate and performance summary across CLOSED trades."""
    closed = [t for t in trades if t["status"] != "OPEN"]
    wins = [t for t in closed if t["status"] in ("TP1_HIT", "TP2_HIT")]
    losses = [t for t in closed if t["status"] == "SL_HIT"]
    manual = [t for t in closed if t["status"] == "CLOSED_MANUAL"]
    win_rate = (len(wins) / len(closed) * 100) if closed else 0
    avg_win_pnl = sum(t["pnl_pct"] for t in wins if t.get("pnl_pct") is not None) / len(wins) if wins else 0
    avg_loss_pnl = sum(t["pnl_pct"] for t in losses if t.get("pnl_pct") is not None) / len(losses) if losses else 0
    open_count = len([t for t in trades if t["status"] == "OPEN"])
    return {
        "total": len(trades), "open": open_count, "closed": len(closed),
        "wins": len(wins), "losses": len(losses), "manual_closes": len(manual),
        "win_rate": round(win_rate, 1),
        "avg_win_pnl": round(avg_win_pnl, 2), "avg_loss_pnl": round(avg_loss_pnl, 2),
    }


# ─────────────────────────────────────────────────────────────
#   LIVE SYMBOL LISTS (for the no-screenshot live dashboard)
# ─────────────────────────────────────────────────────────────
def get_spot_symbols():
    """All live Bitget spot pairs, e.g. BTCUSDT, ETHUSDT..."""
    try:
        r = requests.get("https://api.bitget.com/api/v2/spot/public/symbols", timeout=10)
        data = r.json().get("data", [])
        out = []
        for d in data:
            if d.get("status") == "online" and d.get("quoteCoin") == "USDT":
                out.append({
                    "symbol": d.get("symbol"),
                    "base": d.get("baseCoin"),
                    "quote": d.get("quoteCoin"),
                })
        return sorted(out, key=lambda x: x["base"])
    except Exception:
        return []


def get_futures_symbols():
    """All live Bitget futures contracts across every product category
    (crypto USDT-M, USDC-M, coin-margined, and stock-linked like AAPL/TSLA).
    Deduped by BASE COIN, not just by exact symbol string - the same coin
    (e.g. BTC) can be listed under several product categories with different
    symbol names (BTCUSDT, BTCUSD, BTCPERP...), which was making it show up
    2-3 times in the picker. We keep the first one found, and
    FUTURES_PRODUCT_TYPES is ordered so usdt-futures (the standard, most
    liquid contract type) wins that tie-break."""
    out = []
    seen_symbols = set()
    seen_bases = set()
    for ptype in FUTURES_PRODUCT_TYPES:
        try:
            r = requests.get(
                "https://api.bitget.com/api/v2/mix/market/contracts",
                params={"productType": ptype}, timeout=10,
            )
            data = r.json().get("data", [])
            for d in data:
                sym = d.get("symbol")
                base = d.get("baseCoin")
                if not sym or sym in seen_symbols:
                    continue
                if not base or base in seen_bases:
                    continue
                if d.get("symbolStatus") == "normal":
                    out.append({
                        "symbol": sym,
                        "base": base,
                        "quote": d.get("quoteCoin"),
                        "product_type": ptype,
                    })
                    seen_symbols.add(sym)
                    seen_bases.add(base)
        except Exception:
            continue
    return sorted(out, key=lambda x: x["base"] or "")


def get_all_tickers(market_type="spot"):
    """Batch live price/24h-change for ALL symbols in one call — used for the
    fast browse/filter table so we don't hit the API once per coin."""
    out = {}
    try:
        if market_type == "spot":
            r = requests.get("https://api.bitget.com/api/v2/spot/market/tickers", timeout=12)
            data = r.json().get("data", [])
            for d in data:
                sym = d.get("symbol")
                if not sym:
                    continue
                try:
                    price = float(d.get("lastPr") or d.get("last") or 0)
                    chg = float(d.get("change24h") or d.get("changeUtc24h") or 0) * 100
                    vol = float(d.get("usdtVolume") or d.get("baseVolume") or 0)
                except Exception:
                    price, chg, vol = 0, 0, 0
                out[sym] = {"price": price, "change_24h": chg, "volume": vol}
        else:
            for ptype in FUTURES_PRODUCT_TYPES:
                try:
                    r = requests.get(
                        "https://api.bitget.com/api/v2/mix/market/tickers",
                        params={"productType": ptype}, timeout=12,
                    )
                    data = r.json().get("data", [])
                    for d in data:
                        sym = d.get("symbol")
                        if not sym or sym in out:
                            continue
                        try:
                            price = float(d.get("lastPr") or d.get("last") or 0)
                            chg = float(d.get("change24h") or d.get("changeUtc24h") or 0) * 100
                            vol = float(d.get("usdtVolume") or d.get("baseVolume") or 0)
                        except Exception:
                            price, chg, vol = 0, 0, 0
                        out[sym] = {"price": price, "change_24h": chg, "volume": vol}
                except Exception:
                    continue
        return out
    except Exception:
        return out


def build_auto_chart(coin_symbol, pair, market_type, timeframe, live_price, indicators):
    """Builds the same 'chart' dict the Gemini vision step used to produce —
    but purely from live Bitget data + calculated indicators. No screenshot,
    no AI vision call needed for the main live dashboard."""
    ema9 = indicators.get("ema9")
    ema21 = indicators.get("ema21")
    ema50 = indicators.get("ema50")
    price = live_price or 0

    if ema9 and ema21 and ema50:
        if price > ema9 > ema21 > ema50:
            trend = "Uptrend"
        elif price < ema9 < ema21 < ema50:
            trend = "Downtrend"
        elif price > ema21:
            trend = "Uptrend"
        elif price < ema21:
            trend = "Downtrend"
        else:
            trend = "Sideways"
    else:
        trend = "Sideways"

    swing_sup = indicators.get("swing_support")
    swing_res = indicators.get("swing_resistance")

    coin_id = COIN_MAP.get(coin_symbol, coin_symbol.lower())

    # chart_confidence used to be hardcoded "HIGH" on every single live
    # analysis regardless of what data actually came back — meaning it
    # always contributed the same fixed bonus everywhere. Now it reflects
    # how complete the indicator set really was for this coin/timeframe.
    completeness_checks = [ema50, indicators.get("ema200"), indicators.get("atr"),
                           indicators.get("bb_upper"), indicators.get("bb_lower"),
                           swing_sup is not None and swing_res is not None]
    complete_count = sum(1 for v in completeness_checks if v)
    if complete_count >= 5:
        data_confidence = "HIGH"
    elif complete_count >= 3:
        data_confidence = "MODERATE"
    else:
        data_confidence = "LOW"

    return {
        "coin_symbol": coin_symbol, "pair": pair, "coin_id": coin_id,
        "timeframe": timeframe, "mkt_type": "Futures" if market_type == "futures" else "Spot",
        "price": price, "trend": trend,
        "support": f"{swing_sup:.6f}" if swing_sup else "N/A",
        "resistance": f"{swing_res:.6f}" if swing_res else "N/A",
        "volume": indicators.get("vol_signal", "N/A"),
        "buyer_seller": "N/A", "ma_sig": "N/A",
        "chart_confidence": data_confidence,
        "reason": "Live data-based analysis (RSI, MACD, EMA, Bollinger Bands, ATR, swing levels, order book, funding, sentiment).",
        "warning": "None",
    }


# ─────────────────────────────────────────────────────────────
#   GEMINI
# ─────────────────────────────────────────────────────────────
def get_gemini_response(prompt, image, api_key, log=None):
    client = genai.Client(api_key=api_key)
    for model_name in GEMINI_MODELS:
        try:
            response = client.models.generate_content(model=model_name, contents=[prompt, image])
            return response
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                if log: log(f"⚠️ {model_name} quota exceeded, trying next model...")
                time.sleep(2)
                continue
            else:
                if log: log(f"Model error ({model_name}): {e}")
                continue
    if log: log("All Gemini models are over quota right now.")
    return None


# ─────────────────────────────────────────────────────────────
#   PATTERN LIBRARY (stored as JSON file next to the app)
# ─────────────────────────────────────────────────────────────
def load_library(github_token=None):
    if github_token:
        content, _ = _github_get("data/pattern_library.json", github_token)
        if content is not None:
            try:
                library = json.loads(content)
                with open(PATTERN_LIBRARY_FILE, "w") as f:
                    f.write(content)
                return library
            except Exception:
                pass
    if os.path.exists(PATTERN_LIBRARY_FILE):
        with open(PATTERN_LIBRARY_FILE, "r") as f:
            return json.load(f)
    return {}


def save_library(library, github_token=None):
    content = json.dumps(library, indent=2)
    with open(PATTERN_LIBRARY_FILE, "w") as f:
        f.write(content)
    if github_token:
        _github_put("data/pattern_library.json", github_token, content, "Update pattern library")


def process_pattern_response(raw, filename, library):
    added_names = []
    for block in raw.split("---"):
        block = block.strip()
        if not block:
            continue
        pat = {}
        for line in block.split("\n"):
            if ":" in line:
                key, _, val = line.partition(":")
                pat[key.strip()] = val.strip()
        name = pat.get("PATTERN_NAME", "").strip()
        if not name:
            continue
        key = name.upper().replace(" ", "_")
        library[key] = {
            "name": name,
            "type": pat.get("TYPE", "Unknown"),
            "signal": pat.get("SIGNAL", "Unknown"),
            "description": pat.get("DESCRIPTION", ""),
            "conditions": pat.get("CONDITIONS", ""),
            "entry": pat.get("ENTRY", ""),
            "target": pat.get("TARGET", ""),
            "stop_loss": pat.get("STOP_LOSS", ""),
            "reliability": pat.get("RELIABILITY", "Medium"),
            "added_from": filename,
            "added_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
        added_names.append(name)
    return added_names


def add_pattern_from_image(image, filename, api_key, library):
    prompt = """You are a technical analysis expert.
This image shows candlestick chart patterns.

Respond in this EXACT format for EACH pattern:

PATTERN_NAME: [name]
TYPE: [Bullish / Bearish / Neutral]
SIGNAL: [Reversal / Continuation / Indecision]
DESCRIPTION: [2-3 lines]
CONDITIONS: [required market conditions]
ENTRY: [when to enter]
TARGET: [expected move]
STOP_LOSS: [where to place SL]
RELIABILITY: [High / Medium / Low]
---"""
    response = get_gemini_response(prompt, image, api_key)
    if not response:
        return []
    return process_pattern_response(response.text.strip(), filename, library)


def add_patterns_from_pdf(pdf_bytes, filename, api_key, library, progress_cb=None):
    import fitz
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    added_total = []
    for i, page in enumerate(pdf):
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        added = add_pattern_from_image(img, filename, api_key, library)
        added_total.extend(added)
        if progress_cb:
            progress_cb(i + 1, len(pdf))
    pdf.close()
    return added_total


# ─────────────────────────────────────────────────────────────
#   CHART READING
# ─────────────────────────────────────────────────────────────
def analyze_chart_full(image, api_key):
    prompt = """Expert crypto technical analyst. Read this chart ONLY for visual data.
Extract ONLY what you can SEE — do not guess or suggest trades.

COIN_SYMBOL: [visible symbol]
PAIR: [visible pair]
TIMEFRAME: [visible timeframe]
MARKET_TYPE: [Spot or Futures]
CURRENT_PRICE: [exact price on chart]
TREND: [Uptrend / Downtrend / Sideways]
SUPPORT: [visible support level]
RESISTANCE: [visible resistance level]
VOLUME: [Increasing / Decreasing / Low / High]
BUYER_SELLER: [if order book visible, else N/A]
MA_SIGNAL: [Price above MAs / Price below MAs / Mixed / N/A]
CHART_CONFIDENCE: [HIGH / MODERATE / LOW — how clearly readable is this chart]
REASON: [3-4 lines of what you visually see on chart]
WARNING: [any visible red flags]"""

    response = get_gemini_response(prompt, image, api_key)
    if not response:
        return None
    raw = response.text.strip()
    result = {}
    for line in raw.split("\n"):
        if ":" in line:
            key, _, val = line.partition(":")
            result[key.strip()] = val.strip()

    def get(k, d=""):
        return result.get(k, d).strip()

    def to_float(s):
        try:
            return float(re.sub(r"[^0-9.]", "", s.split()[0]))
        except Exception:
            return None

    coin_symbol = get("COIN_SYMBOL", "UNKNOWN").upper().replace("/", "").replace("USDT", "")
    pair = get("PAIR", coin_symbol + "USDT").upper().replace(" ", "")
    coin_id = COIN_MAP.get(coin_symbol, coin_symbol.lower())

    return {
        "coin_symbol": coin_symbol, "pair": pair, "coin_id": coin_id,
        "timeframe": get("TIMEFRAME", "?"), "mkt_type": get("MARKET_TYPE", "Futures"),
        "price": to_float(get("CURRENT_PRICE", "0")),
        "trend": get("TREND", "?"), "support": get("SUPPORT", "N/A"),
        "resistance": get("RESISTANCE", "N/A"), "volume": get("VOLUME", "N/A"),
        "buyer_seller": get("BUYER_SELLER", "N/A"), "ma_sig": get("MA_SIGNAL", "N/A"),
        "chart_confidence": get("CHART_CONFIDENCE", "MODERATE"),
        "reason": get("REASON", "N/A"), "warning": get("WARNING", "None"),
    }


def match_patterns_from_chart(image, library, api_key):
    if not library:
        return [], "Neutral"
    pattern_list = "\n".join([f"- {p['name']} ({p['type']}): {p['description']}" for p in library.values()])
    prompt = f"""Expert technical analyst. Check this chart for these patterns:
{pattern_list}

Respond EXACTLY:
PATTERNS_FOUND: [names or NONE]
PATTERN_SIGNAL: [Bullish/Bearish/Neutral]
PATTERN_CONFIDENCE: [HIGH/MODERATE/LOW]"""

    response = get_gemini_response(prompt, image, api_key)
    if not response:
        return [], "Neutral"
    raw = response.text.strip()
    matched = []
    overall_signal = "Neutral"
    for line in raw.split("\n"):
        line = line.strip()
        if line.startswith("PATTERNS_FOUND:"):
            found = line.split(":", 1)[1].strip()
            if found.upper() != "NONE" and found:
                for n in [x.strip() for x in found.split(",")]:
                    for key, pat in library.items():
                        if pat["name"].lower() in n.lower() or n.lower() in pat["name"].lower():
                            matched.append(pat)
                            break
        elif line.startswith("PATTERN_SIGNAL:"):
            overall_signal = line.split(":", 1)[1].strip()
    return matched, overall_signal


# ─────────────────────────────────────────────────────────────
#   LIVE MARKET DATA
# ─────────────────────────────────────────────────────────────
def get_market_data(coin_id, coin_symbol):
    try:
        r = requests.get(
            f"https://api.coingecko.com/api/v3/coins/{coin_id}",
            params={"localization": "false", "market_data": "true", "community_data": "true"},
            timeout=12,
        )
        d = r.json()
        md = d.get("market_data", {})
        return {
            "price": md.get("current_price", {}).get("usd", 0),
            "mkt_cap": md.get("market_cap", {}).get("usd", 0),
            "vol_24h": md.get("total_volume", {}).get("usd", 0),
            "ch_1h": md.get("price_change_percentage_1h_in_currency", {}).get("usd", 0) or 0,
            "ch_24h": md.get("price_change_percentage_24h", 0) or 0,
            "ch_7d": md.get("price_change_percentage_7d", 0) or 0,
            "sent_up": d.get("sentiment_votes_up_percentage", 50) or 50,
        }
    except Exception:
        return {}


def get_orderbook(pair, market_type="spot"):
    try:
        if market_type == "futures":
            for ptype in FUTURES_PRODUCT_TYPES:
                try:
                    r = requests.get(
                        "https://api.bitget.com/api/v2/mix/market/merge-depth",
                        params={"symbol": pair, "productType": ptype, "limit": "20"}, timeout=10,
                    )
                    data = r.json().get("data", {})
                    asks = data.get("asks", [])
                    bids = data.get("bids", [])
                    if asks and bids:
                        ta = sum(float(a[1]) for a in asks)
                        tb = sum(float(b[1]) for b in bids)
                        t = ta + tb
                        return {"buy_pct": (tb / t) * 100, "sell_pct": (ta / t) * 100}
                except Exception:
                    continue
            return {"buy_pct": 50, "sell_pct": 50}
        else:
            r = requests.get(
                "https://api.bitget.com/api/v2/spot/market/orderbook",
                params={"symbol": pair, "limit": "20"}, timeout=10,
            )
            data = r.json().get("data", {})
            asks = data.get("asks", [])
            bids = data.get("bids", [])
            if not asks or not bids:
                raise Exception("empty")
            ta = sum(float(a[1]) for a in asks)
            tb = sum(float(b[1]) for b in bids)
            t = ta + tb
            return {"buy_pct": (tb / t) * 100, "sell_pct": (ta / t) * 100}
    except Exception:
        return {"buy_pct": 50, "sell_pct": 50}


def get_fear_greed():
    try:
        r = requests.get("https://api.alternative.me/fng/?limit=1", timeout=8)
        d = r.json()["data"][0]
        return {"value": int(d["value"]), "label": d["value_classification"]}
    except Exception:
        return {"value": 50, "label": "Neutral"}


def get_funding_rate(pair, market_type="futures"):
    if market_type != "futures":
        return {"rate": 0, "signal": "NEUTRAL"}
    for ptype in FUTURES_PRODUCT_TYPES:
        try:
            url = "https://api.bitget.com/api/v2/mix/market/current-fund-rate"
            r = requests.get(url, params={"symbol": pair, "productType": ptype}, timeout=10)
            data = r.json().get("data", [])
            if not data:
                continue
            rate = float(data[0].get("fundingRate", 0)) * 100
            if rate > 0.1:
                signal = "SHORT"
            elif rate > 0.05:
                signal = "SHORT"
            elif rate < -0.1:
                signal = "LONG"
            elif rate < -0.05:
                signal = "LONG"
            else:
                signal = "NEUTRAL"
            return {"rate": rate, "signal": signal}
        except Exception:
            continue
    return {"rate": 0, "signal": "NEUTRAL"}


def get_realtime_indicators(pair, timeframe="1h", market_type="spot"):
    try:
        # IMPORTANT: Bitget's Spot candles endpoint and Futures/Mix candles
        # endpoint use DIFFERENT casing for the same granularity. Spot wants
        # lowercase ("1h", "4h", "1day"), Futures/Mix wants uppercase
        # ("1H", "4H", "1D"). Using the wrong case returns an empty/invalid
        # response, not an error - which is why timeframes like 2h/4h/1d were
        # silently failing on futures while 5m/15m/30m (same either way) worked.
        tf_map_spot = {"1m": "1min", "3m": "3min", "5m": "5min", "15m": "15min", "30m": "30min",
                       "1h": "1h", "2h": "2h", "4h": "4h", "6h": "6h", "12h": "12h",
                       "1d": "1day", "1w": "1week"}
        tf_map_futures = {"1m": "1m", "3m": "3m", "5m": "5m", "15m": "15m", "30m": "30m",
                          "1h": "1H", "2h": "2H", "4h": "4H", "6h": "6H", "12h": "12H",
                          "1d": "1D", "1w": "1W"}
        if market_type == "futures":
            tf = tf_map_futures.get(timeframe.lower(), "1H")
        else:
            tf = tf_map_spot.get(timeframe.lower(), "1h")
        candles = []
        if market_type == "futures":
            # Some listings (e.g. stock-linked futures like AAPL/TSLA) live under a
            # different productType than plain crypto perpetuals — try all of them.
            for ptype in FUTURES_PRODUCT_TYPES:
                try:
                    r = requests.get(
                        "https://api.bitget.com/api/v2/mix/market/candles",
                        params={"symbol": pair, "granularity": tf, "limit": "200", "productType": ptype},
                        timeout=10,
                    )
                    candles = r.json().get("data", [])
                    if candles:
                        break
                except Exception:
                    continue
        else:
            r = requests.get(
                "https://api.bitget.com/api/v2/spot/market/candles",
                params={"symbol": pair, "granularity": tf, "limit": "200"}, timeout=10,
            )
            candles = r.json().get("data", [])
        if not candles:
            return {}
        candles.reverse()
        closes = [float(c[4]) for c in candles]
        highs = [float(c[2]) for c in candles]
        lows = [float(c[3]) for c in candles]
        volumes = [float(c[5]) for c in candles]

        def rsi(prices, p=14):
            g, l = [], []
            for i in range(1, len(prices)):
                d = prices[i] - prices[i - 1]
                g.append(max(d, 0))
                l.append(max(-d, 0))
            if len(g) < p:
                return 50
            ag = sum(g[-p:]) / p
            al = sum(l[-p:]) / p
            return 50 if al == 0 else round(100 - (100 / (1 + ag / al)), 2)

        def ema(prices, p):
            if len(prices) < p:
                return None
            k = 2 / (p + 1)
            e = sum(prices[:p]) / p
            for x in prices[p:]:
                e = x * k + e * (1 - k)
            return round(e, 6)

        def macd(prices):
            e12 = ema(prices, 12)
            e26 = ema(prices, 26)
            if not e12 or not e26:
                return None, None, None
            ml = round(e12 - e26, 6)
            mv = [ema(prices[:i + 1], 12) - ema(prices[:i + 1], 26) for i in range(26, len(prices))
                  if ema(prices[:i + 1], 12) and ema(prices[:i + 1], 26)]
            sig = ema(mv, 9) if len(mv) >= 9 else None
            return ml, round(sig, 6) if sig else None, round(ml - sig, 6) if sig else None

        def bb(prices, p=20, s=2):
            if len(prices) < p:
                return None, None, None
            r_ = prices[-p:]
            m = sum(r_) / p
            std = (sum((x - m) ** 2 for x in r_) / p) ** 0.5
            return round(m + s * std, 6), round(m, 6), round(m - s * std, 6)

        def stoch_rsi(prices, p=14):
            rv = [rsi(prices[max(0, i - p):i + 1], p) for i in range(p, len(prices))]
            if len(rv) < p:
                return 50
            rc = rv[-p:]
            mn, mx = min(rc), max(rc)
            return 50 if mx == mn else round((rv[-1] - mn) / (mx - mn) * 100, 2)

        def vol_sig(vols):
            if len(vols) < 20:
                return "Unknown"
            avg = sum(vols[-20:-1]) / 19
            c = vols[-1]
            if c > avg * 1.5:
                return "High (Strong move likely)"
            elif c > avg * 1.2:
                return "Above Average"
            elif c < avg * 0.7:
                return "Low (Weak move)"
            return "Average"

        def atr_calc(h, l, c, p=14):
            if len(c) < p + 1:
                return None
            trs = []
            for i in range(1, len(c)):
                tr = max(h[i] - l[i], abs(h[i] - c[i - 1]), abs(l[i] - c[i - 1]))
                trs.append(tr)
            if len(trs) < p:
                return None
            return round(sum(trs[-p:]) / p, 8)

        def swing_levels(h, l, c, lookback=60):
            n = min(lookback, len(h))
            rh = h[-n:]
            rl = l[-n:]
            current = c[-1]
            resistances, supports = [], []
            for i in range(2, len(rh) - 2):
                if rh[i] > rh[i - 1] and rh[i] > rh[i - 2] and rh[i] > rh[i + 1] and rh[i] > rh[i + 2]:
                    resistances.append(rh[i])
                if rl[i] < rl[i - 1] and rl[i] < rl[i - 2] and rl[i] < rl[i + 1] and rl[i] < rl[i + 2]:
                    supports.append(rl[i])
            res_above = [x for x in resistances if x > current]
            sup_below = [x for x in supports if x < current]
            return (max(sup_below) if sup_below else None,
                    min(res_above) if res_above else None)

        RSI = rsi(closes)
        EMA9 = ema(closes, 9)
        EMA21 = ema(closes, 21)
        EMA50 = ema(closes, 50)
        EMA200 = ema(closes, 200)
        MACD, SIG, HIST = macd(closes)
        BBU, BBM, BBL = bb(closes)
        SRSI = stoch_rsi(closes)
        VSIG = vol_sig(volumes)
        CP = closes[-1]
        ATR = atr_calc(highs, lows, closes)
        SWING_SUP, SWING_RES = swing_levels(highs, lows, closes)

        rt = "LONG" if RSI <= 30 or RSI >= 55 and RSI < 70 else ("SHORT" if RSI >= 70 or RSI <= 45 else "NEUTRAL")
        if RSI >= 70:
            rt = "SHORT"
        elif RSI <= 30:
            rt = "LONG"
        elif RSI >= 55:
            rt = "LONG"
        elif RSI <= 45:
            rt = "SHORT"
        else:
            rt = "NEUTRAL"

        if MACD and SIG:
            if MACD > SIG and HIST and HIST > 0:
                mt = "LONG"
            elif MACD < SIG and HIST and HIST < 0:
                mt = "SHORT"
            else:
                mt = "NEUTRAL"
        else:
            mt = "NEUTRAL"

        if EMA9 and EMA21 and EMA50:
            if CP > EMA9 > EMA21 > EMA50:
                et = "LONG"
            elif CP < EMA9 < EMA21 < EMA50:
                et = "SHORT"
            elif CP > EMA21:
                et = "LONG"
            elif CP < EMA21:
                et = "SHORT"
            else:
                et = "NEUTRAL"
        else:
            et = "NEUTRAL"

        if BBU and BBL:
            if CP >= BBU:
                bt = "SHORT"
            elif CP <= BBL:
                bt = "LONG"
            elif CP > BBM:
                bt = "LONG"
            else:
                bt = "SHORT"
        else:
            bt = "NEUTRAL"

        if SRSI >= 80:
            st_ = "SHORT"
        elif SRSI <= 20:
            st_ = "LONG"
        else:
            st_ = "NEUTRAL"

        trades = [rt, mt, et, bt, st_]
        lc = trades.count("LONG")
        sc = trades.count("SHORT")
        idir = "LONG" if lc > sc else ("SHORT" if sc > lc else "NEUTRAL")

        return {
            "rsi": RSI, "stoch_rsi": SRSI, "macd": MACD, "macd_signal": SIG, "histogram": HIST,
            "ema9": EMA9, "ema21": EMA21, "ema50": EMA50, "ema200": EMA200,
            "bb_upper": BBU, "bb_mid": BBM, "bb_lower": BBL,
            "vol_signal": VSIG, "ind_direction": idir, "long_count": lc, "short_count": sc,
            "atr": ATR, "swing_support": SWING_SUP, "swing_resistance": SWING_RES,
            "last_close": CP,
        }
    except Exception:
        return {}


def get_news(coin_symbol, coin_id, newsapi_key):
    """News sentiment via NewsAPI headlines + descriptions (TextBlob).
    IMPORTANT: a missing/blank key, an API error, or too few articles is
    now reported explicitly via 'reliable'/'reason' instead of silently
    coming back as score=0 ('Neutral') - that was misleading, since it
    looked like sentiment was actually checked and happened to be neutral
    when really nothing was checked at all."""
    if not newsapi_key:
        return {"score": 0, "articles": [], "count": 0, "reliable": False,
                "reason": "No NewsAPI key configured"}
    try:
        r = requests.get(
            "https://newsapi.org/v2/everything",
            params={
                "q": f'"{coin_symbol}" OR "{coin_id}"', "language": "en", "sortBy": "publishedAt",
                "pageSize": 15, "apiKey": newsapi_key,
                "from": (datetime.now() - timedelta(days=3)).strftime("%Y-%m-%d"),
            }, timeout=10,
        )
        data = r.json()
        if data.get("status") != "ok":
            return {"score": 0, "articles": [], "count": 0, "reliable": False,
                    "reason": f"NewsAPI error: {data.get('message', 'unknown')}"}
        articles = data.get("articles", [])
        if not articles:
            return {"score": 0, "articles": [], "count": 0, "reliable": False,
                    "reason": "No recent articles found"}
        scored = []
        scores = []
        for a in articles[:12]:
            title = (a.get("title") or "")[:100]
            desc = (a.get("description") or "")[:200]
            text = f"{title}. {desc}".strip()
            pol = TextBlob(text).sentiment.polarity
            scored.append({"title": title, "polarity": round(pol, 3)})
            scores.append(pol)
        avg = sum(scores) / len(scores) if scores else 0
        return {
            "score": avg, "articles": scored, "count": len(scores),
            "reliable": len(scores) >= 3,
            "reason": "OK" if len(scores) >= 3 else "Too few articles for a reliable score",
        }
    except Exception as e:
        return {"score": 0, "articles": [], "count": 0, "reliable": False,
                "reason": f"NewsAPI request failed: {e}"}


# ─────────────────────────────────────────────────────────────
#   HIGHER-TIMEFRAME TREND FILTER
# ─────────────────────────────────────────────────────────────
# Most retail trades lose money by fighting the bigger trend (e.g. taking a
# 15-minute "oversold bounce" LONG while the 4-hour trend is strongly down).
# This checks a higher timeframe than the one selected and reports whether
# it agrees, so final_verdict can flag/penalize counter-trend setups instead
# of only ever looking at one timeframe in isolation.
HTF_MAP = {
    "1m": "1h", "3m": "1h", "5m": "4h", "15m": "4h", "30m": "4h",
    "1h": "1d", "2h": "1d", "4h": "1w", "6h": "1w", "12h": "1w",
    "1d": "1w", "1w": "1w",
}


def get_htf_trend(pair, market_type, timeframe):
    htf_tf = HTF_MAP.get(timeframe, "1d")
    if htf_tf == timeframe:
        return {"available": False, "trend": "NEUTRAL", "timeframe": htf_tf}
    ind = get_realtime_indicators(pair, htf_tf, market_type)
    if not ind:
        return {"available": False, "trend": "NEUTRAL", "timeframe": htf_tf}
    ema50 = ind.get("ema50")
    ema200 = ind.get("ema200")
    cp = ind.get("last_close")
    trend = "NEUTRAL"
    if ema50 and ema200 and cp:
        if cp > ema50 > ema200:
            trend = "LONG"
        elif cp < ema50 < ema200:
            trend = "SHORT"
        elif cp > ema50:
            trend = "LONG"
        elif cp < ema50:
            trend = "SHORT"
    elif ema50 and cp:
        trend = "LONG" if cp > ema50 else "SHORT"
    return {"available": True, "trend": trend, "timeframe": htf_tf}


# ─────────────────────────────────────────────────────────────
#   FINAL VERDICT (data-driven decision + trade levels)
# ─────────────────────────────────────────────────────────────
def final_verdict(chart, market, orderbook, fg, funding, indicators, news, matched_patterns,
                   has_ai_opinion=True, htf=None, target_style="auto"):
    buy_pct = orderbook.get("buy_pct", 50)
    sell_pct = orderbook.get("sell_pct", 50)
    fg_val = fg.get("value", 50)
    news_score = news.get("score", 0)
    news_reliable = news.get("reliable", False)
    community_sent = market.get("sent_up", 50)  # CoinGecko community poll - free, no key needed
    ch_24h = market.get("ch_24h", 0)
    price = chart.get("price") or market.get("price", 0)
    fund_rate = funding.get("rate", 0)
    fund_signal = funding.get("signal", "NEUTRAL")
    rsi_val = indicators.get("rsi", 50)
    ind_dir = indicators.get("ind_direction", "NEUTRAL")
    long_c = indicators.get("long_count", 0)
    short_c = indicators.get("short_count", 0)
    confidence = chart.get("chart_confidence", "MODERATE")
    atr = indicators.get("atr")
    swing_sup = indicators.get("swing_support")
    swing_res = indicators.get("swing_resistance")
    bb_lower = indicators.get("bb_lower")
    bb_upper_v = indicators.get("bb_upper")
    htf = htf or {"available": False, "trend": "NEUTRAL", "timeframe": "-"}

    # ── Directional vote — decides direction ONLY ──────────────────────
    ls = 0
    ss = 0
    if buy_pct > 55: ls += 1
    elif sell_pct > 55: ss += 1
    if fg_val < 45: ls += 1
    elif fg_val > 65: ss += 1
    if news_reliable:
        if news_score > 0.1: ls += 1
        elif news_score < -0.1: ss += 1
    else:
        # Not enough real news articles to trust NewsAPI's number this time
        # - fall back to CoinGecko's free community-sentiment poll (already
        # fetched, no extra key needed) instead of silently voting neutral.
        if community_sent >= 60: ls += 1
        elif community_sent <= 40: ss += 1
    if fund_signal == "LONG": ls += 2
    elif fund_signal == "SHORT": ss += 2
    if ind_dir == "LONG": ls += 3
    elif ind_dir == "SHORT": ss += 3
    if ch_24h > 1: ls += 1
    elif ch_24h < -1: ss += 1
    if htf.get("trend") == "LONG": ls += 2
    elif htf.get("trend") == "SHORT": ss += 2

    margin = ls - ss
    max_margin = 11  # sum of all weights above: 1+1+1+2+3+1+2
    # Require a real margin, not just any lead - a 1-point win among 7
    # weighted categories used to be enough to call a full trade direction.
    if margin >= 2:
        data_direction = "LONG"
    elif margin <= -2:
        data_direction = "SHORT"
    else:
        data_direction = "NEUTRAL"

    if not has_ai_opinion:
        final_direction = data_direction if data_direction != "NEUTRAL" else (
            "LONG" if ch_24h >= 0 else "SHORT"
        )
        agreement = "FULL" if data_direction != "NEUTRAL" else "PARTIAL"
        gemini_direction = final_direction
    else:
        trend_lower = chart.get("trend", "").lower()
        if "up" in trend_lower:
            gemini_direction = "LONG"
        elif "down" in trend_lower:
            gemini_direction = "SHORT"
        else:
            gemini_direction = "LONG" if ch_24h > 0 else "SHORT"

        if data_direction == gemini_direction:
            agreement = "FULL"
            final_direction = data_direction
        elif data_direction == "NEUTRAL":
            agreement = "PARTIAL"
            final_direction = gemini_direction
        else:
            agreement = "CONFLICT"
            final_direction = data_direction

    # ── Confidence score — INDEPENDENT measures, not a re-check of the
    # same votes used above. Re-scoring the same signals that already
    # decided the direction was circular: whichever direction won would
    # automatically "confirm itself" and score high, regardless of whether
    # it was actually a good trade. Confidence now comes from: how decisive
    # the vote margin was, how many of the 5 technical indicators actually
    # agree with each other, whether the higher timeframe trend agrees, and
    # how complete the underlying data was.
    factors = []

    if confidence == "HIGH":
        factors.append(("good", "Full indicator set available (EMA/BB/ATR/swing all calculated)"))
    elif confidence == "MODERATE":
        factors.append(("warn", "Partial indicator set — some values missing"))
    else:
        factors.append(("bad", "Indicator data mostly incomplete"))

    if matched_patterns:
        bull_p = [p for p in matched_patterns if "bull" in p.get("type", "").lower()]
        bear_p = [p for p in matched_patterns if "bear" in p.get("type", "").lower()]
        if final_direction == "LONG" and bull_p:
            factors.append(("good", f"Bullish patterns matched: {', '.join([p['name'] for p in bull_p])}"))
        elif final_direction == "SHORT" and bear_p:
            factors.append(("good", f"Bearish patterns matched: {', '.join([p['name'] for p in bear_p])}"))
        else:
            factors.append(("warn", "Patterns found but don't clearly support this direction"))
    else:
        factors.append(("warn", "No chart patterns detected"))

    winning_votes = ls if final_direction == "LONG" else ss
    losing_votes = ss if final_direction == "LONG" else ls
    if margin >= 4:
        factors.append(("good", f"Strong directional vote ({winning_votes} vs {losing_votes} across all signals)"))
    elif margin >= 2:
        factors.append(("warn", f"Moderate directional vote ({winning_votes} vs {losing_votes})"))
    else:
        factors.append(("bad", f"Thin/no clear majority ({winning_votes} vs {losing_votes}) — low-conviction setup"))

    ind_winning = long_c if final_direction == "LONG" else short_c
    if ind_winning >= 4:
        factors.append(("good", f"Indicators strongly agree ({ind_winning}/5 point this way)"))
    elif ind_winning == 3:
        factors.append(("warn", f"Indicators lean this way ({ind_winning}/5)"))
    else:
        factors.append(("bad", f"Indicators are split ({ind_winning}/5) — no clear technical edge"))

    if htf.get("available"):
        if htf["trend"] == final_direction:
            factors.append(("good", f"Higher timeframe ({htf['timeframe']}) trend agrees"))
        elif htf["trend"] == "NEUTRAL":
            factors.append(("warn", f"Higher timeframe ({htf['timeframe']}) trend is flat/unclear"))
        else:
            factors.append(("bad", f"⚠️ Against higher timeframe ({htf['timeframe']}) trend — counter-trend, higher risk"))
    else:
        factors.append(("warn", "Higher-timeframe trend check unavailable"))

    if fund_signal != "NEUTRAL":
        aligned = (fund_signal == final_direction)
        factors.append(("good" if aligned else "warn",
                         f"Funding rate {fund_rate:+.4f}% {'supports' if aligned else 'does not support'} this trade"))
    else:
        factors.append(("warn", f"Funding rate neutral ({fund_rate:+.4f}%)"))

    if news_reliable:
        matches = (news_score > 0.1) == (final_direction == "LONG")
        factors.append(("good" if matches else "warn",
                         f"News sentiment: {news_score:+.2f} from {news.get('count', 0)} articles"))
    else:
        factors.append(("warn", f"News: {news.get('reason', 'unavailable')} — used CoinGecko community sentiment ({community_sent:.0f}% bullish) as backup"))

    if fg_val <= 25 or fg_val >= 75:
        factors.append(("warn", f"Fear & Greed at an extreme ({fg_val}, {fg.get('label','')}) — contrarian risk"))
    else:
        factors.append(("good", f"Fear & Greed in a normal range ({fg_val}, {fg.get('label','')})"))

    margin_pct = min(abs(margin) / max_margin, 1.0) * 100
    ind_pct = (ind_winning / 5) * 100
    if htf.get("available"):
        htf_pct = 100 if htf["trend"] == final_direction else (50 if htf["trend"] == "NEUTRAL" else 0)
    else:
        htf_pct = 50
    data_pct = {"HIGH": 100, "MODERATE": 60}.get(confidence, 30)

    accuracy = round(margin_pct * 0.40 + ind_pct * 0.25 + htf_pct * 0.20 + data_pct * 0.15, 1)
    accuracy = max(0, min(accuracy, 100))

    # ── Trade levels — anchored to the ACTUAL planned entry price ──────
    # Previously SL/TP/R:R were calculated off today's live price even
    # though the entry itself was set above/below that price (waiting for
    # breakout confirmation). That meant the real risk once filled at the
    # confirmation price was bigger than shown, and the real reward was
    # smaller than shown. Everything below is now anchored to entry_ref,
    # the actual price the trader will enter at.
    tp1 = tp2 = sl = entry_low = entry_high = None
    entry_note = ""
    confirm_price = invalidate_price = None
    rev_confirm = rev_entry_low = rev_entry_high = rev_tp1 = rev_tp2 = rev_sl = None
    rr = "N/A"

    # TP distance is now decided BY the analysis itself, not a manual
    # choice. A weak/thin setup gets a closer TP (bank profit before the
    # low-conviction read has time to be wrong); a setup with a strong
    # vote margin AND a higher-timeframe trend actually agreeing gets more
    # room, since there's real evidence the move can run further. This
    # directly uses the same independent strength measures as the
    # confidence score, so "how far is TP" and "how confident is this"
    # always tell a consistent story instead of a fixed distance being
    # slapped on regardless of setup quality.
    if target_style == "auto":
        if accuracy >= 75 and htf_pct >= 50:
            eff_style = "aggressive"
        elif accuracy >= 55:
            eff_style = "balanced"
        else:
            eff_style = "conservative"
    else:
        eff_style = target_style

    tp_mult = {
        "conservative": (1.5, 2.5),
        "balanced": (2.0, 3.5),
        "aggressive": (2.5, 4.5),
    }.get(eff_style, (2.0, 3.5))
    tp1_mult, tp2_mult = tp_mult
    factors.append(("warn" if eff_style == "conservative" else "good",
                     f"Target distance: {eff_style} ({tp1_mult}x/{tp2_mult}x risk) — auto-picked from this setup's own strength"))

    if price and price > 0:
        eff_atr = atr if (atr and atr > 0) else price * 0.01
        ctx = []
        if final_direction == "LONG":
            if rsi_val <= 35: ctx.append("RSI oversold")
            if bb_lower and price <= bb_lower * 1.01: ctx.append("near BB lower band")
            if fund_signal == "LONG": ctx.append("funding favors longs")
        else:
            if rsi_val >= 65: ctx.append("RSI overbought")
            if bb_upper_v and price >= bb_upper_v * 0.99: ctx.append("near BB upper band")
            if fund_signal == "SHORT": ctx.append("funding favors shorts")
        ctx_txt = (" — " + ", ".join(ctx)) if ctx else ""

        if final_direction == "LONG":
            entry_ref = round(price + eff_atr * 0.35, 8)
            if swing_sup and swing_sup < price and (price - swing_sup) < eff_atr * 6:
                sl = round(swing_sup - eff_atr * 0.3, 8)
            else:
                sl = round(price - eff_atr * 1.5, 8)
            risk = entry_ref - sl
            if swing_res and swing_res > entry_ref and (swing_res - entry_ref) >= risk * 1.3:
                tp1 = round(swing_res, 8)
            else:
                tp1 = round(entry_ref + risk * tp1_mult, 8)
            tp2 = round(entry_ref + risk * tp2_mult, 8)
            confirm_price = entry_ref
            invalidate_price = round(sl - eff_atr * 0.15, 8)
            entry_low = confirm_price
            entry_high = round(confirm_price + eff_atr * 0.25, 8)
            entry_note = f"Enter after candle confirms above this level{ctx_txt}"
            rev_confirm = invalidate_price
            rev_entry_low = round(rev_confirm - eff_atr * 0.25, 8)
            rev_entry_high = rev_confirm
            rev_risk = eff_atr * 1.5
            rev_tp1 = round(rev_confirm - rev_risk * tp1_mult, 8)
            rev_tp2 = round(rev_confirm - rev_risk * tp2_mult, 8)
            rev_sl = round(entry_ref + eff_atr * 0.5, 8)
        else:
            entry_ref = round(price - eff_atr * 0.35, 8)
            if swing_res and swing_res > price and (swing_res - price) < eff_atr * 6:
                sl = round(swing_res + eff_atr * 0.3, 8)
            else:
                sl = round(price + eff_atr * 1.5, 8)
            risk = sl - entry_ref
            if swing_sup and swing_sup < entry_ref and (entry_ref - swing_sup) >= risk * 1.3:
                tp1 = round(swing_sup, 8)
            else:
                tp1 = round(entry_ref - risk * tp1_mult, 8)
            tp2 = round(entry_ref - risk * tp2_mult, 8)
            confirm_price = entry_ref
            invalidate_price = round(sl + eff_atr * 0.15, 8)
            entry_low = round(confirm_price - eff_atr * 0.25, 8)
            entry_high = confirm_price
            entry_note = f"Enter after candle confirms below this level{ctx_txt}"
            rev_confirm = invalidate_price
            rev_entry_low = rev_confirm
            rev_entry_high = round(rev_confirm + eff_atr * 0.25, 8)
            rev_risk = eff_atr * 1.5
            rev_tp1 = round(rev_confirm + rev_risk * tp1_mult, 8)
            rev_tp2 = round(rev_confirm + rev_risk * tp2_mult, 8)
            rev_sl = round(entry_ref - eff_atr * 0.5, 8)

        rr = round(abs(tp1 - entry_ref) / abs(sl - entry_ref), 1) if sl and abs(sl - entry_ref) > 0 else "N/A"

    return {
        "final_direction": final_direction, "gemini_direction": gemini_direction,
        "data_direction": data_direction, "agreement": agreement, "accuracy": accuracy,
        "vote_margin": margin, "htf_trend": htf.get("trend"), "htf_timeframe": htf.get("timeframe"),
        "target_style": eff_style,
        "factors": factors, "tp1": tp1, "tp2": tp2, "sl": sl,
        "entry_low": entry_low, "entry_high": entry_high, "entry_note": entry_note,
        "confirm_price": confirm_price, "invalidate_price": invalidate_price,
        "rev_confirm": rev_confirm, "rev_entry_low": rev_entry_low, "rev_entry_high": rev_entry_high,
        "rev_tp1": rev_tp1, "rev_tp2": rev_tp2, "rev_sl": rev_sl, "rr": rr, "price": price,
    }


# ─────────────────────────────────────────────────────────────
#   WORD REPORT
# ─────────────────────────────────────────────────────────────
def generate_docx_bytes(chart, market, funding, indicators, verdict, matched_patterns, chart_image_pil):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    GREEN = RGBColor(0x00, 0x99, 0x44)
    RED = RGBColor(0xD3, 0x2F, 0x2F)
    YELLOW = RGBColor(0xB2, 0x8A, 0x00)
    NAVY = RGBColor(0x1a, 0x23, 0x7e)
    GRAY = RGBColor(0x55, 0x55, 0x55)

    def set_cell(cell, text, bold=False, color=None, size=10, center=False):
        cell.text = ""
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AI SMART TRADE ANALYZER")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run(f"Trade Decision Report — {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

    if chart_image_pil:
        try:
            img_copy = chart_image_pil.copy()
            if img_copy.mode in ("RGBA", "P"):
                img_copy = img_copy.convert("RGB")
            max_w = 1400
            if img_copy.width > max_w:
                ratio = max_w / img_copy.width
                img_copy = img_copy.resize((max_w, int(img_copy.height * ratio)), Image.LANCZOS)
            buf = io.BytesIO()
            img_copy.save(buf, format="JPEG", quality=78)
            buf.seek(0)
            doc.add_picture(buf, width=Inches(6.3))
        except Exception:
            pass

    doc.add_paragraph()
    h = doc.add_heading("Final Decision", level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY

    accuracy = verdict["accuracy"]
    agreement = verdict["agreement"]
    final_direction = verdict["final_direction"]

    dec_p = doc.add_paragraph()
    dec_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if agreement == "CONFLICT":
        dtxt, dclr = "MARKET NOT IN FAVOUR — SKIP THIS TRADE", RED
    elif accuracy >= 75:
        dtxt, dclr = f"ENTER TRADE — {final_direction}", GREEN
    elif accuracy >= 55:
        dtxt, dclr = f"WAIT FOR CONFIRMATION — {final_direction}", YELLOW
    else:
        dtxt, dclr = "SKIP — WEAK SETUP", RED
    dr = dec_p.add_run(dtxt)
    dr.bold = True
    dr.font.size = Pt(16)
    dr.font.color.rgb = dclr

    doc.add_paragraph()

    if agreement != "CONFLICT" and verdict["tp1"] and verdict["sl"] and verdict["entry_low"]:
        h2 = doc.add_heading("Trade Levels", level=2)
        for run in h2.runs:
            run.font.color.rgb = NAVY
        t2 = doc.add_table(rows=5, cols=3)
        t2.style = "Table Grid"
        rows_data = [
            ("Entry Zone", f"${verdict['entry_low']:,.6f} - ${verdict['entry_high']:,.6f}", verdict["entry_note"]),
            ("Take Profit 1", f"${verdict['tp1']:,.6f}", "Conservative target"),
            ("Take Profit 2", f"${verdict['tp2']:,.6f}", "Extended target"),
            ("Stop Loss", f"${verdict['sl']:,.6f}", "Mandatory - always set this"),
            ("Risk / Reward", f"1 : {verdict['rr']}", ""),
        ]
        colors_map = [NAVY, GREEN, GREEN, RED, NAVY]
        for i, (label, val, note) in enumerate(rows_data):
            set_cell(t2.cell(i, 0), label, bold=True)
            set_cell(t2.cell(i, 1), val, bold=True, color=colors_map[i])
            set_cell(t2.cell(i, 2), note, color=GRAY, size=9)
        doc.add_paragraph()

    h4 = doc.add_heading("Market Overview", level=2)
    for run in h4.runs:
        run.font.color.rgb = NAVY

    def fmt(n):
        if n >= 1_000_000_000: return f"${n/1_000_000_000:.2f}B"
        if n >= 1_000_000: return f"${n/1_000_000:.2f}M"
        return f"${n:,.4f}"

    price = verdict["price"]
    overview_rows = [
        ("Coin / Pair", f"{chart.get('coin_symbol','?')} / {chart.get('pair','?')}", "Market Type", chart.get("mkt_type", "?")),
        ("Timeframe", chart.get("timeframe", "?"), "Trend", chart.get("trend", "?")),
        ("Live Price", f"${price:,.6f}" if price else "N/A", "24h Change", f"{market.get('ch_24h',0):+.2f}%"),
        ("Market Cap", fmt(market.get("mkt_cap", 0)), "24h Volume", fmt(market.get("vol_24h", 0))),
        ("Chart Support", chart.get("support", "N/A"), "Chart Resistance", chart.get("resistance", "N/A")),
        ("Funding Rate", f"{funding.get('rate',0):+.4f}%", "Funding Signal", funding.get("signal", "N/A")),
    ]
    t4 = doc.add_table(rows=len(overview_rows), cols=4)
    t4.style = "Table Grid"
    for i, (a, b, c_, d) in enumerate(overview_rows):
        set_cell(t4.cell(i, 0), a, bold=True, size=9)
        set_cell(t4.cell(i, 1), b, size=9)
        set_cell(t4.cell(i, 2), c_, bold=True, size=9)
        set_cell(t4.cell(i, 3), d, size=9)

    doc.add_paragraph()
    h5 = doc.add_heading("Technical Indicators", level=2)
    for run in h5.runs:
        run.font.color.rgb = NAVY
    atr_val = indicators.get("atr")
    atr_pct = (atr_val / price * 100) if (atr_val and price) else 0
    ind_rows = [
        ("RSI (14)", f"{indicators.get('rsi',0):.2f}", "Stoch RSI", f"{indicators.get('stoch_rsi',0):.2f}"),
        ("EMA 21", f"{indicators.get('ema21',0):.4f}" if indicators.get("ema21") else "N/A",
         "EMA 50", f"{indicators.get('ema50',0):.4f}" if indicators.get("ema50") else "N/A"),
        ("BB Upper", f"{indicators.get('bb_upper',0):.4f}" if indicators.get("bb_upper") else "N/A",
         "BB Lower", f"{indicators.get('bb_lower',0):.4f}" if indicators.get("bb_lower") else "N/A"),
        ("ATR (Volatility)", f"{atr_val:.6f} ({atr_pct:.2f}%)" if atr_val else "N/A",
         "Volume", indicators.get("vol_signal", "N/A")),
    ]
    t5 = doc.add_table(rows=len(ind_rows), cols=4)
    t5.style = "Table Grid"
    for i, (a, b, c_, d) in enumerate(ind_rows):
        set_cell(t5.cell(i, 0), a, bold=True, size=9)
        set_cell(t5.cell(i, 1), b, size=9)
        set_cell(t5.cell(i, 2), c_, bold=True, size=9)
        set_cell(t5.cell(i, 3), d, size=9)

    doc.add_paragraph()
    h7 = doc.add_heading("Signal Breakdown", level=2)
    for run in h7.runs:
        run.font.color.rgb = NAVY
    for level, text in verdict["factors"]:
        clr = GREEN if level == "good" else (RED if level == "bad" else YELLOW)
        fp = doc.add_paragraph(f"• {text}")
        for run in fp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = clr

    doc.add_paragraph()
    fp1 = doc.add_paragraph()
    fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr1 = fp1.add_run("Always use stop-loss. Max 2% risk per trade. This is AI analysis - not financial advice.")
    fr1.italic = True
    fr1.font.size = Pt(8)
    fr1.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────
#   FULL PIPELINE (called by the Streamlit app)
# ─────────────────────────────────────────────────────────────
def run_full_analysis(image, gemini_key, newsapi_key, library, market_type="spot", log=lambda msg: None):
    """Screenshot-based deep-dive (uses Gemini vision) — optional extra mode."""
    log("Reading chart with Gemini...")
    chart = analyze_chart_full(image, gemini_key)
    if not chart:
        return None

    log("Matching known patterns...")
    matched, pat_signal = match_patterns_from_chart(image, library, gemini_key)

    log("Fetching live market data...")
    market = get_market_data(chart["coin_id"], chart["coin_symbol"])

    log("Fetching order book...")
    orderbook = get_orderbook(chart["pair"], market_type)

    log("Fetching Fear & Greed index...")
    fg = get_fear_greed()

    log("Fetching funding rate...")
    funding = get_funding_rate(chart["pair"], market_type)

    log("Calculating technical indicators...")
    indicators = get_realtime_indicators(chart["pair"], chart["timeframe"], market_type)

    log("Checking news sentiment...")
    news = get_news(chart["coin_symbol"], chart["coin_id"], newsapi_key)

    log("Checking higher-timeframe trend...")
    htf = get_htf_trend(chart["pair"], market_type, chart["timeframe"])

    log("Building final verdict...")
    verdict = final_verdict(chart, market, orderbook, fg, funding, indicators, news, matched,
                             has_ai_opinion=True, htf=htf)

    return {
        "chart": chart, "market": market, "orderbook": orderbook, "fg": fg,
        "funding": funding, "indicators": indicators, "news": news,
        "matched_patterns": matched, "verdict": verdict,
    }


def run_live_analysis(coin_symbol, pair, market_type, timeframe, newsapi_key,
                       use_news=True, log=lambda msg: None):
    """No-screenshot pipeline — the main live-dashboard mode. Pulls everything
    straight from Bitget/CoinGecko/Fear&Greed and builds the same verdict
    (Entry / TP / SL / Direction / Confidence) automatically."""
    log(f"Calculating indicators for {coin_symbol}...")
    indicators = get_realtime_indicators(pair, timeframe, market_type)
    if not indicators:
        return {"error": f"No candle data returned for {pair} ({market_type}, {timeframe}). "
                          f"Bitget API might be rate-limiting or the symbol/timeframe combo is unsupported."}

    log("Fetching order book...")
    orderbook = get_orderbook(pair, market_type)

    log("Fetching Fear & Greed index...")
    fg = get_fear_greed()

    log("Fetching funding rate...")
    funding = get_funding_rate(pair, market_type)

    log("Fetching live price from Bitget...")
    # Use Bitget's own ticker for THIS exact pair as the primary price source,
    # so it matches what's on screen on Bitget. CoinGecko is an aggregated,
    # slightly-delayed price across many exchanges and can legitimately
    # differ from Bitget's own last-traded price - it should never be the
    # primary source for a Bitget-based tool.
    live_price = get_single_ticker_price(pair, market_type)

    log("Fetching market cap / 24h data...")
    coin_id = COIN_MAP.get(coin_symbol, coin_symbol.lower())
    market = get_market_data(coin_id, coin_symbol)

    if not live_price:
        # Bitget ticker failed (rare) - fall back to CoinGecko, then to the
        # last closed candle price used for the indicators.
        live_price = market.get("price") or indicators.get("last_close") or 0

    news = {"score": 0, "articles": [], "count": 0, "reliable": False, "reason": "News check disabled"}
    if use_news:
        log("Checking news sentiment...")
        news = get_news(coin_symbol, coin_id, newsapi_key)

    log("Checking higher-timeframe trend...")
    htf = get_htf_trend(pair, market_type, timeframe)

    chart = build_auto_chart(coin_symbol, pair, market_type, timeframe, live_price, indicators)

    log("Building final verdict...")
    verdict = final_verdict(chart, market, orderbook, fg, funding, indicators, news,
                             matched_patterns=[], has_ai_opinion=False, htf=htf)

    return {
        "chart": chart, "market": market, "orderbook": orderbook, "fg": fg,
        "funding": funding, "indicators": indicators, "news": news,
        "matched_patterns": [], "verdict": verdict,
    }
